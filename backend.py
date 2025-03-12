import os
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import time
def delete_folder(folder_path):
    """
    Delete a folder and all its contents.
    """
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        for root, dirs, files in os.walk(folder_path, topdown=False):
            for file in files:
                os.remove(os.path.join(root, file))
            for dir in dirs:
                os.rmdir(os.path.join(root, dir))
        os.rmdir(folder_path)
        print(f"Deleted folder: {folder_path}")
    else:
        print(f"Folder not found: {folder_path}")

def pdf_to_images(pdf_path, output_folder, start_page=1, end_page=None):
    """
    Convert a PDF into images (one per page).
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    pdf_document = fitz.open(pdf_path)
    total_pages = len(pdf_document)
    print(f"Number of pages in the PDF: {total_pages}")

    if end_page is None or end_page > total_pages:
        end_page = total_pages

    extracted_images = []

    for page_number in range(start_page - 1, end_page):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        image_path = os.path.join(output_folder, f"page_{page_number + 1}.jpg")
        pix.save(image_path)
        extracted_images.append(image_path)
        print(f"Saved image: {image_path}")

    return extracted_images

def set_paragraph_direction(paragraph, direction='rtl'):
    """
    Set the text direction of a paragraph to RTL or LTR.
    """
    p_pr = paragraph._element.get_or_add_pPr()
    p_bidi = OxmlElement('w:bidi')
    p_bidi.set(qn('w:val'), '1' if direction == 'rtl' else '0')
    p_pr.append(p_bidi)

def fix_inverted_brackets(text):
    """
    Fix inverted brackets in RTL text by adding Unicode control characters.
    """
    # Characters that should not be mirrored in RTL text
    fixed_text = []
    for char in text:
        if char in '()[]{}':
            # Add Unicode LEFT-TO-RIGHT MARK (LRM) before and after the character
            fixed_text.append('\u200E' + char + '\u200E')
        else:
            fixed_text.append(char)
    return ''.join(fixed_text)

def html_to_docx(html_content, doc):
    """
    Convert HTML content to DOCX format while preserving formatting and alignment.
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['h1', 'p', 'ol', 'li']):
        if element.name == 'h1':
            # Add heading
            heading = doc.add_heading(level=1)
            run = heading.add_run(fix_inverted_brackets(element.get_text()))
            run.font.size = Pt(16)
            if 'style' in element.attrs and 'text-align: center' in element['style']:
                set_paragraph_direction(heading, 'rtl')  # Set RTL for Arabic text
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            else:
                set_paragraph_direction(heading, 'rtl')  # Set RTL for Arabic text

            set_paragraph_direction(heading, 'rtl')  # Set RTL for Arabic text

        elif element.name == 'p':
            # Add paragraph
            paragraph = doc.add_paragraph()
            # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Always align paragraphs to the right
            set_paragraph_direction(paragraph, 'rtl')  # Set RTL for Arabic text

            # Handle bold text
            for content in element.contents:
                if content.name == 'b':
                    run = paragraph.add_run(fix_inverted_brackets(content.get_text()))
                    run.bold = True
                else:
                    paragraph.add_run(fix_inverted_brackets(str(content)))

        elif element.name == 'ol':
    # Add ordered list
            list_number = 1  # Reset list numbering for each new page
            for li in element.find_all('li'):
                # Add a paragraph without the 'List Number' style
                paragraph = doc.add_paragraph()
                set_paragraph_direction(paragraph, 'rtl')  # Set RTL for Arabic text

                # Add the list number manually
                paragraph.add_run(f" .{list_number} ").bold = True
                list_number += 1  # Increment the list number

                # Add the list item content
                for content in li.contents:
                    if content.name == 'b':
                        run = paragraph.add_run(fix_inverted_brackets(content.get_text()))
                        run.bold = True
                    else:
                        paragraph.add_run(fix_inverted_brackets(str(content)))

    # Add a page break after the list
    doc.add_page_break()