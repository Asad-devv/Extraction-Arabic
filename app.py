import os
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import streamlit as st
from dotenv import load_dotenv
import time
from backend import delete_folder,pdf_to_images,html_to_docx
# Load environment variables
load_dotenv()

prompt="""Extract the Arabic text from the provided image and format it in an HTML page. The text should be structured as follows:

1. If the page contains a heading, enclose it in `<h1>`. Center it only if the original document has it centered, using `style="text-align: center;"`.
2. Main paragraphs should be in `<p>` tags, aligned to the right using `style="text-align: right; direction: rtl;"`. Ensure that each paragraph is enclosed in a single `<p>` tag, even if it spans multiple lines in the original document.
3. Bold important words using `<b>`.
4. Any footnotes should be formatted as an ordered list `<ol>` at the end of the page, with each list item `<li>` aligned to the right using `style="text-align: right; direction: rtl;"`.
5. Preserve the original document structure as much as possible.
6. Do not use a `<style>` tag or external CSS. Apply all styles inline using the `style` attribute.
7. Do not break paragraphs into multiple `<p>` tags for each line. Use a single `<p>` tag for the entire paragraph.
8. Do not add any other HTML tags apart from the ones mentioned above (`<h1>`, `<p>`, `<b>`, `<ol>`, and `<li>`).

Return only the HTML content without additional explanations."""

# Function to perform find and replace in a DOCX file
def find_and_replace_in_docx(doc, find_texts, replace_texts):
    if len(find_texts) != len(replace_texts):
        raise ValueError("Find and Replace lists must have the same length.")

    for find_text, replace_text in zip(find_texts, replace_texts):
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find_text in cell.text:
                        cell.text = cell.text.replace(find_text, replace_text)

# Streamlit UI
st.sidebar.header("Navigation")
options = ["Process PDF", "Find and Replace"]
choice = st.sidebar.radio("Go to:", options)

# Process PDF Section
if choice == "Process PDF":
    st.title("Arabic PDF to Word Converter")
    st.write("Upload a PDF, extract Arabic content, and download the result in a Word document.")

    # Input fields
    user_api_key = st.text_input("Enter your Gemini API Key (optional):", type="password")
    pdf_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    start_page = st.number_input("Start Page (1-based index):", value=1, min_value=1)
    end_page = st.number_input("End Page (inclusive):", value=1, min_value=1)
    st.markdown("""
    ***Process PDF Section Notes:***
    If you do not provide an API key:

    The tool will process a maximum of 100 pages starting from the specified start page.
    You can run the tool multiple times (e.g., 60 times) to process more than 60 pages by \nadjusting the start and end pages for each run.
    Example: If your PDF has 300 pages, you can process pages 1-100 in the first run, pages 101-200 in the second run, and pages 201-300 in the third run.

    If you provide an API key:
    The tool will process all pages within the specified range without any limitations.

    Example: If your PDF has 500 pages and you provide an API key, you can process all 500 pages in a single run.""")
    if st.button("Process PDF"):
        if not pdf_file:
            st.error("Please upload a PDF file.")
        else:
            try:
                # Save the uploaded PDF
                pdf_path = os.path.join("temp", "uploaded_pdf.pdf")
                os.makedirs("temp", exist_ok=True)
                with open(pdf_path, "wb") as f:
                    f.write(pdf_file.read())

                # Validate and enforce page limits
                pdf_document = fitz.open(pdf_path)
                total_pages = len(pdf_document)
                pdf_document.close()

                if end_page == 0 or end_page > total_pages:
                    end_page = total_pages

                # Limit to 100 pages if no API key is provided
                if not user_api_key:
                    st.warning("API key not provided. Limiting processing to 100 pages.")
                    if (end_page-start_page>100):
                        end_page = min(start_page + 99, total_pages)
                    else:
                        end_page=end_page

                # Convert PDF pages to images
                output_folder = "temp_images"
                image_paths = pdf_to_images(pdf_path, output_folder, start_page=start_page, end_page=end_page)

                # Initialize Word document
                doc = Document()

                # Extract content and process pages
                st.write("Extracting content from the PDF...")

                # API Keys (Use them in rotation)
                API_KEY = [
                    os.getenv("API_KEY_1"),
                    os.getenv("API_KEY_2"),
                    os.getenv("API_KEY_3"),
                    os.getenv("API_KEY_4"),
                    os.getenv("API_KEY_5"),
                ]

                i = 0
                for index, image_path in enumerate(image_paths, start=start_page):
                    try:
                        print(f"Using API Key: {API_KEY[i]}")
                        genai.configure(api_key=API_KEY[i])
                        i = (i + 1) % len(API_KEY)

                        model = genai.GenerativeModel("models/gemini-2.0-flash")
                        st.write(f"Processing: {image_path}")

                        # Upload the image
                        myfile = genai.upload_file(image_path)

                        # Generate content
                        result = model.generate_content([myfile, prompt])
                        html_content = result.text
                        print(html_content)
                            
                        # Convert HTML to DOCX
                        html_to_docx(html_content, doc)
                        time.sleep(1)

                    except Exception as e:
                        st.error(f"Error processing page {index}: {e}")

                # Save the Word document
                output_file_name = f"{os.path.splitext(pdf_file.name)[0]}_pages_{start_page}_to_{end_page}.docx"
                output_path = os.path.join("temp", output_file_name)
                doc.save(output_path)

                # Provide a download link
                with open(output_path, "rb") as f:
                    st.download_button("Download Word Document", f, file_name=output_file_name)

            except Exception as e:
                st.error(f"Error: {e}")
            
            delete_folder("temp_images")
            delete_folder("temp")
# Find and Replace Section
elif choice == "Find and Replace":
    st.title("Find and Replace in Arabic DOCX")
    st.write("Upload a DOCX file, specify text to find and replace, and download the updated document.")
    st.markdown("""
    The Find and Replace feature is designed for further cleaning and refining of the DOCX file extracted from Gemini. Here’s how to use it:

    Remove Unwanted Text:
        To remove specific text, add the text you want to remove in the Find box and leave the Replace box empty (or add a space).
        Example:
        Find: (Example Text)
        Replace: (a single space)
        This will remove (Example Text) from the document.
        
    Replace Text:
        To replace specific text, add the text you want to replace in the Find box and the new text in the Replace box.
        Example:
        Find: الشيخ الأكبر
        Replace: محيي الدين ابن عربي
        This will replace all occurrences of الشيخ الأكبر with محيي الدين ابن عربي.
        
    Fix Formatting Issues:
        Use Find and Replace to fix formatting issues, such as extra spaces or unwanted characters.
        Example:
        Find: (double space)
        Replace: (single space)
        This will replace all double spaces with single spaces.
        
    Add or Modify Text:
        Use Find and Replace to add or modify specific phrases or words.
        Example:
        Find: ملحق
        Replace: ملحق توضيحي
        This will replace all occurrences of ملحق with ملحق توضيحي.
""")
    # Inject CSS to align text inputs to the right
    st.markdown(
        """
        <style>
        .right-align input {
            text-align: right !important;
        }
        .stTextInput input {
            text-align: right !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    docx_file = st.file_uploader("Upload a DOCX file for Editing", type=["docx"])

    # Initialize session state for dynamic find-replace inputs
    if "find_replace_pairs" not in st.session_state:
        st.session_state.find_replace_pairs = [("", "")]

    st.subheader("Specify Text to Find and Replace (Use copy-paste for quick and better results)")

    # Dynamic inputs for find and replace pairs
    for i, (find_text, replace_text) in enumerate(st.session_state.find_replace_pairs):
        cols = st.columns(2)
        with cols[0]:
            st.session_state.find_replace_pairs[i] = (
                st.text_input(
                    f"Text to Find {i + 1} (Arabic):",
                    value=find_text,
                    key=f"find_{i}",
                    placeholder="Enter text to find",
                ),
                st.session_state.find_replace_pairs[i][1]
            )
        with cols[1]:
            st.session_state.find_replace_pairs[i] = (
                st.session_state.find_replace_pairs[i][0],
                st.text_input(
                    f"Replacement Text {i + 1} (Arabic):",
                    value=replace_text,
                    key=f"replace_{i}",
                    placeholder="Enter replacement text",
                )
            )

    # Button to add another pair of inputs
    if st.button("Add Another Find-Replace Pair"):
        st.session_state.find_replace_pairs.append(("", ""))

    output_file_name_edit = st.text_input("Enter output Word file name (without extension):", "مُتَجَدِّدة يَوْميًّا")
    output_file_name_edit += ".docx"

    if st.button("Perform Find and Replace"):
        if not docx_file:
            st.error("Please upload a DOCX file.")
        else:
            try:
                doc_path = os.path.join("temp", "uploaded_docx.docx")
                os.makedirs("temp", exist_ok=True)
                with open(doc_path, "wb") as f:
                    f.write(docx_file.read())

                doc = Document(doc_path)

                # Filter out empty find-replace pairs
                find_replace_pairs = [
                    (find_text.strip(), replace_text.strip())
                    for find_text, replace_text in st.session_state.find_replace_pairs
                    if find_text.strip()  # Include only valid "find" texts
                ]

                # Perform find and replace
                find_and_replace_in_docx(doc, *zip(*find_replace_pairs))

                # Save the updated document
                updated_path = os.path.join("temp", output_file_name_edit)
                doc.save(updated_path)

                # Provide download link
                with open(updated_path, "rb") as f:
                    st.download_button("Download Updated DOCX", f, file_name=output_file_name_edit)

            except Exception as e:
                st.error(f"Error processing the document: {e}")


